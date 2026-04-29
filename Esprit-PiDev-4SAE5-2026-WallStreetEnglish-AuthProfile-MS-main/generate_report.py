"""
Generates the Bug Report as a Word (.docx) file.
Run: python generate_report.py
Output: Bug_Report_WallStreetEnglish.docx
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB) if level == 1 else RGBColor(0x1E, 0x40, 0xAF)
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F3F4F6")
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table_row(table, col1, col2):
    row = table.add_row()
    row.cells[0].text = col1
    row.cells[1].text = col2

# ── Cover ─────────────────────────────────────────────────────────────────────

title = doc.add_heading("Bug Report", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("WallStreet English — Auth & Profile Microservices")
run.font.size = Pt(13)
run.bold = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Esprit PiDev 4SAE5 2026   |   Date: 2026-04-29").font.size = Pt(10)

doc.add_paragraph()

# ── Executive Summary ─────────────────────────────────────────────────────────

set_heading(doc, "Executive Summary")
add_para(doc,
    "Two critical bugs were identified and resolved in the Auth/Profile microservices. "
    "The first caused a 500 Internal Server Error whenever an admin attempted to create a new user. "
    "The second caused user role changes to not be reflected in the Angular dashboard until the "
    "browser page was manually refreshed. Both issues stemmed from architectural problems in how "
    "the two microservices (auth-service and user-management-service) communicated with each other "
    "and with Keycloak.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# BUG 1
# ══════════════════════════════════════════════════════════════════════════════

set_heading(doc, "Bug #1 — POST /api/users Returns 500 Internal Server Error")

# Symptom
set_heading(doc, "Symptom", 2)
add_para(doc,
    "When an administrator clicked \"Create User\" in the Angular admin dashboard, the HTTP call "
    "POST http://localhost:8089/api/users consistently returned 500 Internal Server Error. "
    "The Angular console logged:")
add_code_block(doc,
    "POST http://localhost:8089/api/users  500 (Internal Server Error)\n"
    "HttpErrorResponse { status: 500, url: 'http://localhost:8089/api/users' }")

# Root cause
set_heading(doc, "Root Cause: Circular Service Dependency", 2)
add_para(doc,
    "The admin create-user flow triggered a hidden circular call between the two microservices. "
    "The sequence was:")
add_code_block(doc,
    "Angular\n"
    "  → POST /api/users  (user-management-service)\n"
    "      → saves UserProfile to DB  [transaction OPEN, not committed]\n"
    "      → calls auth-service POST /api/auth/admin/create\n"
    "          → authService.register()\n"
    "              → creates user in Keycloak ✅\n"
    "              → saves user to auth_db ✅\n"
    "              → calls BACK user-management-service POST /api/users/sync-from-auth\n"
    "                  → tries to INSERT same email again\n"
    "                  → DB UNIQUE constraint violation 💥\n"
    "                    (outer transaction not visible due to MySQL isolation)\n"
    "          → auth-service returns error\n"
    "      → user-management-service deletes the profile\n"
    "      → throws RuntimeException → 500 ❌")

add_para(doc,
    "MySQL's default isolation level (REPEATABLE READ) prevents a new transaction from seeing "
    "uncommitted data from another transaction. So when auth-service called back to create the "
    "profile, the existing (uncommitted) profile was invisible, causing a duplicate key violation.")

# Secondary causes
set_heading(doc, "Secondary Causes", 2)

add_para(doc, "2a — Wrong Feign Return Type", bold=True)
add_para(doc,
    "The Feign client in user-management-service expected AuthResponse, "
    "but AdminController.createUserByAdmin() returned ResponseEntity<Map<String,Object>>. "
    "This mismatch caused Feign deserialization errors.")
add_code_block(doc,
    "// Feign client (user-management-service)\n"
    "AuthResponse createUserByAdmin(...);  // ← expects AuthResponse\n\n"
    "// AdminController (auth-service) — WRONG\n"
    "public ResponseEntity<?> createUserByAdmin(...) {\n"
    "    Map<String, Object> result = new HashMap<>();  // ← returns Map\n"
    "    return ResponseEntity.ok(result);\n"
    "}")

add_para(doc, "2b — @PreAuthorize Blocking Internal Service Calls", bold=True)
add_para(doc,
    "The /api/auth/admin/create endpoint had @PreAuthorize(\"hasRole('ADMIN')\"). "
    "When called service-to-service from user-management-service, role mapping differences "
    "could cause 403 Forbidden, which Feign rethrew as a RuntimeException → 500.")

add_para(doc, "2c — Missing firstName/lastName in AuthRegisterRequest", bold=True)
add_para(doc,
    "The DTO used by the Feign client was missing firstName and lastName fields, "
    "so Keycloak user records were created without the user's name.")

# Fix
set_heading(doc, "Fix Applied", 2)

add_para(doc, "Fix 1 — New registerByAdmin() method in AuthService", bold=True)
add_para(doc,
    "A new method was added that creates the user in Keycloak and auth_db WITHOUT "
    "calling back to user-management-service. This completely breaks the circular dependency.")
add_code_block(doc,
    "// AuthService.java — NEW method\n"
    "public AuthResponse registerByAdmin(RegisterRequest request) {\n"
    "    // Creates in Keycloak\n"
    "    // Saves to auth_db\n"
    "    // ⚠️ Does NOT call userServiceClient.createProfile()\n"
    "    //    user-management-service handles its own profile creation\n"
    "}")

add_para(doc, "Fix 2 — Fixed AdminController.createUserByAdmin()", bold=True)
add_code_block(doc,
    "// BEFORE (broken)\n"
    "@PreAuthorize(\"hasRole('ADMIN')\")\n"
    "public ResponseEntity<?> createUserByAdmin(...) {\n"
    "    AuthResponse r = authService.register(request);  // circular call\n"
    "    Map result = new HashMap();                       // wrong return type\n"
    "    return ResponseEntity.ok(result);\n"
    "}\n\n"
    "// AFTER (fixed)\n"
    "// No @PreAuthorize — security enforced by SecurityConfig (.authenticated())\n"
    "public ResponseEntity<AuthResponse> createUserByAdmin(...) {\n"
    "    AuthResponse r = authService.registerByAdmin(request);  // no callback\n"
    "    return ResponseEntity.ok(r);  // correct return type\n"
    "}")

add_para(doc, "Fix 3 — Added firstName/lastName to AuthRegisterRequest", bold=True)
add_code_block(doc,
    "@Data\n"
    "public class AuthRegisterRequest {\n"
    "    private String email;\n"
    "    private String password;\n"
    "    private String confirmPassword;\n"
    "    private String role;\n"
    "    private String firstName;  // ← ADDED\n"
    "    private String lastName;   // ← ADDED\n"
    "}")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# BUG 2
# ══════════════════════════════════════════════════════════════════════════════

set_heading(doc, "Bug #2 — Role Not Synchronized After Admin Updates User Role")

set_heading(doc, "Symptom", 2)
add_para(doc,
    "When an administrator changed a user's role (e.g. from STUDENT to ADMIN) "
    "through the dashboard:")

steps = [
    "The role was correctly updated in Keycloak and both databases ✅",
    "The user was force-logged out ✅",
    "The user logged back in ✅",
    "The Angular dashboard still showed the OLD role page ❌",
    "Only pressing F5 (page refresh) showed the correct new role page ✅",
]
for s in steps:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s)

set_heading(doc, "Root Cause: Stale Role in Login Response", 2)
add_para(doc,
    "AuthService.login() built the login response by reading the role from the "
    "local auth-service database (user.getRole()). However, the auth-service local DB "
    "was not always updated when a role changed — the sync call from user-management-service "
    "to auth-service sometimes failed silently due to the same security issues described in Bug #1.")
add_code_block(doc,
    "// BEFORE (broken) — AuthService.login()\n"
    "User user = userRepository.findByEmail(request.getEmail());\n"
    "return AuthResponse.builder()\n"
    "    .token(accessToken)\n"
    "    .role(user.getRole())   // ← reads from local auth_db — STALE!\n"
    "    .build();")

add_para(doc,
    "So after re-login, the AuthResponse.role returned the OLD role. Angular read this "
    "and showed the wrong dashboard page.")

add_para(doc, "Why page refresh fixed it:", bold=True)
add_para(doc,
    "On a full browser refresh, Angular re-reads the role directly from the Keycloak JWT "
    "token (which was correctly updated) rather than from the login API response. "
    "The JWT is always up to date because Keycloak was the first thing updated during the role change.")

set_heading(doc, "Fix Applied", 2)
add_para(doc,
    "AuthService.login() now reads the role from the Keycloak JWT token — the single "
    "source of truth — instead of from the potentially stale local database. "
    "A new helper extractRoleFromJwt() was added.")
add_code_block(doc,
    "// AFTER (fixed) — AuthService.login()\n"
    "String accessToken = (String) tokenResponse.get(\"access_token\");\n"
    "Role roleFromJwt = extractRoleFromJwt(accessToken);  // decode JWT payload\n\n"
    "if (roleFromJwt != null && roleFromJwt != user.getRole()) {\n"
    "    // Self-heal: update local DB if it was out of date\n"
    "    user.setRole(roleFromJwt);\n"
    "    userRepository.save(user);\n"
    "}\n\n"
    "return AuthResponse.builder()\n"
    "    .token(accessToken)\n"
    "    .role(roleFromJwt)   // ← always from Keycloak JWT ✅\n"
    "    .build();")

add_para(doc,
    "The extractRoleFromJwt() helper base64-decodes the JWT payload and reads "
    "realm_access.roles — the standard Keycloak claim for realm roles (ADMIN, TUTOR, STUDENT). "
    "It falls back to the DB value if the JWT cannot be decoded.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# FILES CHANGED
# ══════════════════════════════════════════════════════════════════════════════

set_heading(doc, "Files Changed")

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "File"
hdr[1].text = "Change"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ("auth-service/.../AuthService.java",
     "Added registerByAdmin() + extractRoleFromJwt() + fixed login() role source"),
    ("auth-service/.../AdminController.java",
     "Fixed return type to AuthResponse, removed @PreAuthorize, call registerByAdmin()"),
    ("user-management-service/.../AuthRegisterRequest.java",
     "Added firstName and lastName fields"),
    ("user-management-service/.../UserProfileService.java",
     "Pass firstName/lastName when building Feign request"),
]
for f, c in rows:
    add_table_row(table, f, c)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# LESSONS LEARNED
# ══════════════════════════════════════════════════════════════════════════════

set_heading(doc, "Architecture Lessons Learned")

lessons = [
    ("Avoid circular service dependencies",
     "Never have Service A call Service B to create resource X, while Service B "
     "simultaneously calls back Service A to create the same resource. The initiating service "
     "should own its own resource creation."),
    ("JWT is the source of truth for roles",
     "Never rely on a local DB copy of roles that may not be in sync with Keycloak. "
     "Always read security-sensitive data directly from the JWT token."),
    ("Silent exception handling hides bugs",
     "Catching exceptions and only logging warnings (without propagating or alerting) "
     "can hide synchronization failures. Consider monitoring/alerting for these cases."),
    ("Internal service endpoints need different security than user-facing ones",
     "Using @PreAuthorize on service-to-service endpoints can block legitimate internal calls. "
     "Use network-level security (private Docker network) for internal endpoints instead."),
]

for title_text, body in lessons:
    add_para(doc, f"• {title_text}", bold=True)
    add_para(doc, f"  {body}")

# ── Save ──────────────────────────────────────────────────────────────────────

output_path = "Bug_Report_WallStreetEnglish.docx"
doc.save(output_path)
print(f"✅ Report saved to: {output_path}")
